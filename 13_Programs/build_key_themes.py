#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import logging
import re
import sys
import uuid
from copy import deepcopy
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

SCHEMA_NAME = "phase2_key_themes_output"
SCHEMA_VERSION = "2.0"
SECTION_NAME = "key_themes"
BOOK_TITLE = "Trump Chronicles"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the Trump Chronicles Key Themes back-matter chapter.")
    parser.add_argument("--week", type=int, required=True, help="Starting week number.")
    parser.add_argument("--weeks", type=int, default=53, help="Number of consecutive weeks to process.")
    parser.add_argument("--force", action="store_true", help="Overwrite existing outputs.")
    parser.add_argument("--dry-run", action="store_true", help="Show planned work without writing files.")
    parser.add_argument(
        "--step3-root",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Step 3/Weeks"),
        help="Root folder containing Week N subfolders.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/10_BackMatter"),
        help="Directory for key_themes outputs.",
    )
    parser.add_argument(
        "--themes-file",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/14_Data/key_themes_payload_v2.json"),
        help="Canonical JSON payload defining the themes and recurring forms.",
    )
    parser.add_argument("--created-by", type=str, default="build_key_themes.py")
    parser.add_argument("--prompt-version", type=str, default="deterministic-v2")
    parser.add_argument("--model", type=str, default=None)
    parser.add_argument("--git-commit", type=str, default=None)
    parser.add_argument(
        "--log-level",
        "--level",
        dest="log_level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="Logging level.",
    )
    return parser.parse_args()


def configure_logging(level: str) -> None:
    logging.basicConfig(level=getattr(logging, level.upper()), format="%(asctime)s [%(levelname)s] %(message)s")


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        payload = json.load(fh)
    if not isinstance(payload, dict):
        raise ValueError(f"{path}: expected JSON object")
    logging.debug("Loaded %s with keys: %s", path, ", ".join(sorted(payload.keys())))
    return payload


def week_numbers(start: int, count: int) -> List[int]:
    if start < 1:
        raise ValueError("--week must be >= 1")
    if count < 1:
        raise ValueError("--weeks must be >= 1")
    return list(range(start, start + count))


def analytic_path(root: Path, week: int) -> Path:
    return root / f"Week {week}" / f"weekly_analytic_brief_week{week}.json"


def metadata_path(root: Path, week: int) -> Path:
    return root / f"Week {week}" / f"weekly_digest_metadata_stack_week{week}.json"


def spine_path(root: Path, week: int) -> Path:
    return root / f"Week {week}" / f"weekly_digest_spine_week{week}.json"


def ensure_writable(output_dir: Path, force: bool) -> Dict[str, Path]:
    json_path = output_dir / "key_themes.json"
    md_path = output_dir / "key_themes.md"
    docx_path = output_dir / "key_themes.docx"
    existing = [p for p in (json_path, md_path, docx_path) if p.exists()]
    if existing and not force:
        raise FileExistsError(f"Output file(s) already exist: {', '.join(str(p) for p in existing)}. Re-run with --force.")
    return {"json": json_path, "md": md_path, "docx": docx_path}


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        text = value
    else:
        text = json.dumps(value, ensure_ascii=False)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def collect_text_fragments(obj: Any) -> List[str]:
    out: List[str] = []
    if obj is None:
        return out
    if isinstance(obj, str):
        cleaned = normalize_text(obj)
        if cleaned:
            out.append(cleaned)
        return out
    if isinstance(obj, list):
        for item in obj:
            out.extend(collect_text_fragments(item))
        return out
    if isinstance(obj, dict):
        for value in obj.values():
            out.extend(collect_text_fragments(value))
        return out
    return out


def extract_week_title(metadata: Dict[str, Any], spine: Dict[str, Any], week: int) -> str:
    title = metadata.get("title") or metadata.get("substack_metadata", {}).get("title") or spine.get("article_title") or f"Week {week}"
    title = normalize_text(title)
    prefix = "This Week in Democracy:"
    if title.startswith(prefix):
        title = title[len(prefix):].strip()
    return title


def week_source_bundle(step3_root: Path, week: int) -> Dict[str, Any]:
    apath = analytic_path(step3_root, week)
    mpath = metadata_path(step3_root, week)
    spath = spine_path(step3_root, week)

    if not apath.exists():
        raise FileNotFoundError(f"Missing analytic brief for Week {week}: {apath}")
    if not mpath.exists():
        raise FileNotFoundError(f"Missing metadata stack for Week {week}: {mpath}")
    if not spath.exists():
        raise FileNotFoundError(f"Missing weekly spine for Week {week}: {spath}")

    analytic = load_json(apath)
    metadata = load_json(mpath)
    spine = load_json(spath)

    return {
        "week": week,
        "analytic_path": str(apath),
        "metadata_path": str(mpath),
        "spine_path": str(spath),
        "analytic": analytic,
        "metadata": metadata,
        "spine": spine,
        "title": extract_week_title(metadata, spine, week),
    }


def build_search_texts(bundle: Dict[str, Any]) -> Dict[str, str]:
    analytic = bundle["analytic"]
    metadata = bundle["metadata"]
    spine = bundle["spine"]

    analytic_parts: List[str] = []
    analytic_parts.extend(collect_text_fragments(analytic.get("summary")))
    analytic_parts.extend(collect_text_fragments(analytic.get("top_themes")))
    analytic_parts.extend(collect_text_fragments(analytic.get("institutional_shifts")))

    metadata_parts: List[str] = []
    metadata_parts.extend(collect_text_fragments(metadata.get("title")))
    metadata_parts.extend(collect_text_fragments(metadata.get("subtitle")))
    metadata_parts.extend(collect_text_fragments(metadata.get("week_summary")))
    metadata_parts.extend(collect_text_fragments(metadata.get("category_flags")))
    metadata_parts.extend(collect_text_fragments(metadata.get("internal_tags")))
    metadata_parts.extend(collect_text_fragments(metadata.get("whole_week_arcs")))

    spine_parts: List[str] = []
    spine_parts.extend(collect_text_fragments(spine.get("article_title")))
    spine_parts.extend(collect_text_fragments(spine.get("article_subtitle")))
    spine_parts.extend(collect_text_fragments(spine.get("sections")))

    return {
        "analytic": " ".join(analytic_parts),
        "metadata": " ".join(metadata_parts),
        "spine": " ".join(spine_parts),
    }


def load_theme_payload(themes_file: Path) -> Dict[str, Any]:
    payload = load_json(themes_file)
    if payload.get("schema_name") != SCHEMA_NAME:
        raise ValueError(f"{themes_file}: schema_name must be {SCHEMA_NAME}")
    if not isinstance(payload.get("themes"), list) or not payload["themes"]:
        raise ValueError(f"{themes_file}: themes array missing or empty")
    return payload


def score_form(patterns: Sequence[str], search_texts: Dict[str, str]) -> Tuple[int, List[Tuple[str, str]]]:
    weights = {"analytic": 4, "metadata": 2, "spine": 3}
    score = 0
    matches: List[Tuple[str, str]] = []
    for pattern in patterns:
        try:
            regex = re.compile(pattern, flags=re.IGNORECASE)
        except re.error:
            regex = re.compile(re.escape(pattern), flags=re.IGNORECASE)
        for source_name, text in search_texts.items():
            if not text:
                continue
            if regex.search(text):
                score += weights[source_name]
                matches.append((source_name, pattern))
    return score, matches


def relevance_from_theme_score(total_score: int, matched_forms: int) -> Optional[str]:
    if matched_forms >= 2 and total_score >= 16:
        return "primary"
    if matched_forms >= 2 and total_score >= 10:
        return "secondary"
    if matched_forms >= 1 and total_score >= 5:
        return "supporting"
    return None


def build_theme_hits(payload: Dict[str, Any], bundles: Sequence[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    theme_hits: Dict[str, Dict[str, Any]] = {}
    for theme in payload["themes"]:
        theme_hits[theme["theme_id"]] = {
            "definition": deepcopy(theme),
            "locator_weeks": [],
            "analytic_brief_weeks": [],
            "metadata_stack_weeks": [],
            "weekly_spine_weeks": [],
            "notes": [],
        }

    for bundle in bundles:
        week = bundle["week"]
        search_texts = build_search_texts(bundle)
        for theme in payload["themes"]:
            form_hits: List[Dict[str, Any]] = []
            total_score = 0
            analytic_hit = False
            metadata_hit = False
            spine_hit = False
            for form in theme.get("recurring_forms", []):
                patterns = form.get("patterns", [])
                if not patterns:
                    continue
                form_score, matches = score_form(patterns, search_texts)
                if form_score <= 0:
                    continue
                total_score += form_score
                sources = sorted({src for src, _ in matches})
                analytic_hit = analytic_hit or ("analytic" in sources)
                metadata_hit = metadata_hit or ("metadata" in sources)
                spine_hit = spine_hit or ("spine" in sources)
                form_hits.append({
                    "form_id": form["form_id"],
                    "form_name": form["form_name"],
                    "score": form_score,
                    "sources": sources,
                    "patterns": sorted({pat for _, pat in matches}),
                })

            matched_forms = len(form_hits)
            relevance = relevance_from_theme_score(total_score, matched_forms)
            if relevance is None:
                continue

            # Penalize excessively broad weak matches so themes do not become universalized.
            if matched_forms == 1 and total_score < 8:
                continue

            theme_hit = theme_hits[theme["theme_id"]]
            theme_hit["locator_weeks"].append({"week": week, "relevance": relevance})
            if analytic_hit:
                theme_hit["analytic_brief_weeks"].append(week)
            if metadata_hit:
                theme_hit["metadata_stack_weeks"].append(week)
            if spine_hit:
                theme_hit["weekly_spine_weeks"].append(week)

            top_forms = sorted(form_hits, key=lambda x: (-x["score"], x["form_name"]))[:3]
            note = "; ".join(
                f"{fh['form_name']} ({fh['score']}, {', '.join(fh['sources'])})" for fh in top_forms
            )
            theme_hit["notes"].append(f"Week {week}: {relevance}; {note}")
            logging.debug(
                "Week %s matched %s: total_score=%s matched_forms=%s relevance=%s",
                week,
                theme["theme_name"],
                total_score,
                matched_forms,
                relevance,
            )
    return theme_hits


def sort_locator_key(locator: Dict[str, Any]) -> Tuple[int, int]:
    relevance_rank = {"primary": 0, "secondary": 1, "supporting": 2}
    return (locator["week"], relevance_rank.get(locator["relevance"], 9))


def build_theme_entries(payload: Dict[str, Any], theme_hits: Dict[str, Dict[str, Any]]) -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []
    for theme in sorted(payload["themes"], key=lambda x: x["sort_order"]):
        hit = theme_hits[theme["theme_id"]]
        if not hit["locator_weeks"]:
            continue
        entry = deepcopy(theme)
        entry["locator_weeks"] = sorted(hit["locator_weeks"], key=sort_locator_key)
        entry["source_summary"] = {
            "analytic_brief_weeks": sorted(set(hit["analytic_brief_weeks"])),
            "metadata_stack_weeks": sorted(set(hit["metadata_stack_weeks"])),
            "weekly_spine_weeks": sorted(set(hit["weekly_spine_weeks"])),
        }
        entry["notes"] = hit["notes"][:40]
        entries.append(entry)
    return entries


def build_output(entries: List[Dict[str, Any]], payload: Dict[str, Any], themes_file: Path, created_by: str, prompt_version: str, model: Optional[str], git_commit: Optional[str]) -> Dict[str, Any]:
    now_utc = datetime.now(UTC)
    run_id = f"key-themes-{now_utc.strftime('%Y%m%dT%H%M%SZ')}-{uuid.uuid4().hex[:8]}"
    generated_from = deepcopy(payload.get("generated_from", {}))
    generated_from["themes_file"] = str(themes_file)
    return {
        "schema_name": SCHEMA_NAME,
        "schema_version": SCHEMA_VERSION,
        "section": SECTION_NAME,
        "book_title": BOOK_TITLE,
        "build": {
            "created_at": now_utc.replace(microsecond=0).isoformat().replace("+00:00", "Z"),
            "created_by": created_by,
            "model": model,
            "run_id": run_id,
            "prompt_version": prompt_version,
            "git_commit": git_commit,
        },
        "generated_from": generated_from,
        "themes": entries,
    }


def collapse_week_ranges(weeks: Sequence[int]) -> str:
    if not weeks:
        return ""
    ordered = sorted(set(weeks))
    ranges: List[str] = []
    start = prev = ordered[0]
    for week in ordered[1:]:
        if week == prev + 1:
            prev = week
            continue
        ranges.append(f"Week {start}" if start == prev else f"Weeks {start}–{prev}")
        start = prev = week
    ranges.append(f"Week {start}" if start == prev else f"Weeks {start}–{prev}")
    return ", ".join(ranges)


def render_markdown(payload: Dict[str, Any]) -> str:
    lines: List[str] = ["# Key Themes", ""]
    for theme in payload["themes"]:
        lines.append(f"## {theme['theme_name']}")
        lines.append("")
        lines.append(theme["framing_line"])
        lines.append("")
        lines.append(theme["body"])
        lines.append("")
        forms_line = "; ".join(f"{form['form_name']}" for form in theme.get("recurring_forms", []))
        if forms_line:
            lines.append(f"**Recurring forms:** {forms_line}.")
            lines.append("")
        locator_text = collapse_week_ranges([loc["week"] for loc in theme["locator_weeks"]])
        lines.append(f"**Prominent in:** {locator_text}")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def render_docx(markdown_text: str, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")
    doc = Document()
    for block in markdown_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            p = doc.add_paragraph()
            p.style = "Heading 1"
            p.add_run(block[2:].strip())
            continue
        if block.startswith("## "):
            p = doc.add_paragraph()
            p.style = "Heading 2"
            p.add_run(block[3:].strip())
            continue
        if block.startswith("**Recurring forms:**"):
            p = doc.add_paragraph()
            run1 = p.add_run("Recurring forms: ")
            run1.bold = True
            p.add_run(block[len("**Recurring forms:**"):].strip())
            continue
        if block.startswith("**Prominent in:**"):
            p = doc.add_paragraph()
            run1 = p.add_run("Prominent in: ")
            run1.bold = True
            p.add_run(block[len("**Prominent in:**"):].strip())
            continue
        doc.add_paragraph(block)
    doc.save(str(output_path))


def main() -> int:
    args = parse_args()
    configure_logging(args.log_level)

    weeks = week_numbers(args.week, args.weeks)
    outputs = {
        "json": args.output_dir / "key_themes.json",
        "md": args.output_dir / "key_themes.md",
        "docx": args.output_dir / "key_themes.docx",
    }

    logging.info("Building Key Themes for Weeks %s to %s", weeks[0], weeks[-1])
    logging.info("Step 3 root: %s", args.step3_root)
    logging.info("Themes file: %s", args.themes_file)
    logging.info("Output dir: %s", args.output_dir)

    if args.dry_run:
        for week in weeks:
            logging.info("[DRY RUN] Would read: %s", analytic_path(args.step3_root, week))
            logging.info("[DRY RUN] Would read: %s", metadata_path(args.step3_root, week))
            logging.info("[DRY RUN] Would read: %s", spine_path(args.step3_root, week))
        logging.info("[DRY RUN] Would read themes payload: %s", args.themes_file)
        for path in outputs.values():
            logging.info("[DRY RUN] Would write: %s", path)
        return 0

    try:
        args.output_dir.mkdir(parents=True, exist_ok=True)
        ensure_writable(args.output_dir, args.force)
        payload = load_theme_payload(args.themes_file)
    except Exception as exc:
        logging.error(str(exc))
        return 1

    bundles: List[Dict[str, Any]] = []
    failed_weeks: List[int] = []
    for week in weeks:
        try:
            bundles.append(week_source_bundle(args.step3_root, week))
            logging.info("Loaded Week %s", week)
        except Exception as exc:
            logging.error("Failed Week %s: %s", week, exc)
            failed_weeks.append(week)

    if not bundles:
        logging.error("No valid weeks loaded; nothing to write.")
        return 1

    theme_hits = build_theme_hits(payload, bundles)
    entries = build_theme_entries(payload, theme_hits)
    if not entries:
        logging.error("No themes identified; nothing to write.")
        return 1

    output = build_output(entries, payload, args.themes_file, args.created_by, args.prompt_version, args.model, args.git_commit)
    markdown = render_markdown(output)

    try:
        with outputs["json"].open("w", encoding="utf-8") as fh:
            json.dump(output, fh, indent=2, ensure_ascii=False)
            fh.write("\n")
        with outputs["md"].open("w", encoding="utf-8") as fh:
            fh.write(markdown)
        render_docx(markdown, outputs["docx"])
    except Exception as exc:
        logging.error("Failed writing outputs: %s", exc)
        return 1

    logging.info("WROTE: %s", outputs["json"])
    logging.info("WROTE: %s", outputs["md"])
    logging.info("WROTE: %s", outputs["docx"])
    if failed_weeks:
        logging.warning("Skipped failed weeks: %s", ", ".join(map(str, failed_weeks)))
    return 0


if __name__ == "__main__":
    sys.exit(main())
