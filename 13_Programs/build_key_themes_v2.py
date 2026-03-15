#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import logging
import re
import sys
import uuid
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


SCHEMA_NAME = "phase2_key_themes_output"
SCHEMA_VERSION = "1.0"
SECTION_NAME = "key_themes"
BOOK_TITLE = "Trump Chronicles"

# Theme locator discipline constants
RELEVANCE_RANK = {"primary": 0, "secondary": 1, "supporting": 2}
PRIMARY_MIN_SCORE = 18
PRIMARY_MIN_PATTERNS = 2
PRIMARY_MIN_SOURCES = 2
SECONDARY_MIN_SCORE = 10
SECONDARY_MIN_PATTERNS = 2
SUPPORTING_MIN_SCORE = 6
SUPPORTING_MIN_PATTERNS = 1
TARGET_PRIMARY_WEEKS = 8
TARGET_TOTAL_WEEKS = 12
FALLBACK_TOTAL_WEEKS = 8
PRIMARY_MIN_FORMS = 2
SECONDARY_MIN_FORMS = 1
SUPPORTING_MIN_FORMS = 1

# Theme-specific rules overrides
THEME_RULE_OVERRIDES: Dict[str, Dict[str, int]] = {
    "THEME_EXECUTIVE_OVERREACH": {
        "primary_min_score": 24,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 14,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_CIVIL_SERVICE_PURGE": {
        "primary_min_score": 20,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 12,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_ATTACKS_ON_DISSENT": {
        "primary_min_score": 20,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 12,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_SELECTIVE_ENFORCEMENT": {
        "primary_min_score": 20,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 12,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_IMMIGRATION_AS_POWER_TEST": {
        "primary_min_score": 22,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 14,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_TRANSACTIONAL_FOREIGN_POLICY": {
        "primary_min_score": 18,
        "primary_min_patterns": 2,
        "primary_min_sources": 2,
        "secondary_min_score": 12,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_COURTS_AS_COUNTERWEIGHT": {
        "primary_min_score": 20,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 12,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 6,
        "target_total_weeks": 8,
        "fallback_total_weeks": 6,
    },
    "THEME_NORMALIZATION_THROUGH_REPETITION": {
        "primary_min_score": 22,
        "primary_min_patterns": 3,
        "primary_min_sources": 2,
        "secondary_min_score": 14,
        "secondary_min_patterns": 2,
        "supporting_min_score": 8,
        "supporting_min_patterns": 2,
        "target_primary_weeks": 5,
        "target_total_weeks": 7,
        "fallback_total_weeks": 5,
    },
}

THEME_FORM_GROUPS: Dict[str, Sequence[Tuple[str, Sequence[str]]]] = {
    "THEME_EXECUTIVE_OVERREACH": (
        ("oversight_bypass", (r"appropriation", r"impound", r"funding freeze", r"circumvent(?:ing)? congress")),
        ("emergency_claims", (r"emergency powers?", r"article ii", r"executive power")),
        ("unilateral_action", (r"unilateral order", r"executive order", r"proclamation")),
    ),
    "THEME_CIVIL_SERVICE_PURGE": (
        ("personnel_purge", (r"mass firings?", r"purges?", r"administrative leave", r"career officials?")),
        ("classification_redesign", (r"schedule [fg]", r"reclassification", r"probationary")),
        ("capacity_hollowing", (r"buyout", r"deferred resignation", r"hiring freeze")),
    ),
    "THEME_ATTACKS_ON_DISSENT": (
        ("campus_pressure", (r"campus protest", r"student visa", r"detained student", r"antisemitism investigation")),
        ("media_pressure", (r"public broadcasting", r"media outlet", r"press credentials?")),
        ("street_protest_pressure", (r"protest crackdown", r"no kings")),
    ),
    "THEME_SELECTIVE_ENFORCEMENT": (
        ("retaliatory_enforcement", (r"retaliat", r"punish perceived opponents", r"make an example")),
        ("double_standard", (r"double standard", r"favored treatment", r"exempt(?:ion|ed) allies")),
    ),
    "THEME_IMMIGRATION_AS_POWER_TEST": (
        ("detention_and_removal", (r"ice", r"cbp", r"deport", r"detention", r"expedited removal")),
        ("exceptional_authority", (r"alien enemies act", r"guant[aá]namo")),
        ("status_and_membership", (r"birthright citizenship", r"asylum")),
    ),
    "THEME_TRANSACTIONAL_FOREIGN_POLICY": (
        ("trade_and_sanctions", (r"tariffs?", r"sanctions?")),
        ("aid_and_security", (r"foreign aid", r"usaid", r"security assistance")),
        ("territorial_or_geopolitical_pressure", (r"gaza", r"panama", r"greenland")),
        ("multilateral_disruption", (r"icc", r"multilateral")),
    ),
    "THEME_COURTS_AS_COUNTERWEIGHT": (
        ("trial_court_blocks", (r"injunction", r"temporary block", r"restraining order", r"preliminary injunction")),
        ("appeals_and_supreme_court", (r"appeals court", r"supreme court")),
        ("litigation_pressure", (r"lawsuit", r"litigation")),
    ),
    "THEME_NORMALIZATION_THROUGH_REPETITION": (
        ("continued_escalation", (r"continued", r"deepened", r"escalated")),
        ("repeat_action", (r"expanded again", r"renewed crackdown", r"another round")),
        ("ordinary_language_of_abuse", (r"normalized", r"routine")),
    ),
}

@dataclass(frozen=True)
class ThemeDef:
    theme_id: str
    theme_name: str
    sort_order: int
    framing_line: str
    body: str
    patterns: Sequence[str] = field(default_factory=tuple)


THEME_CATALOGUE: List[ThemeDef] = [
    ThemeDef(
        "THEME_EXECUTIVE_OVERREACH",
        "Executive Overreach",
        1,
        "The recurrent use of unilateral executive action to outrun deliberation, oversight, and normal institutional process.",
        "This book repeatedly shows the presidency pressing beyond ordinary democratic restraint: using orders, emergency powers, agency control, and unilateral directives to reset policy before Congress, the courts, or the public can fully respond.",
        (r"executive power", r"emergency powers?", r"appropriation", r"impound", r"funding freeze", r"article ii", r"unilateral order", r"circumvent(?:ing)? congress"),
    ),
    ThemeDef(
        "THEME_EXECUTIVE_OLIGARCH_FUSION",
        "Executive–Private Fusion",
        2,
        "The embedding of private actors, business interests, or loyalist networks inside the machinery of the state.",
        "Again and again, the year shows the boundary between public authority and private leverage eroding. This matters democratically because state power becomes harder to audit, less accountable to law, and more vulnerable to patronage, self-dealing, and coercive informality.",
        (r"doge", r"musk", r"private actors?", r"private network", r"payment rails?", r"hr databases?", r"unvetted", r"outsourc", r"privatiz"),
    ),
    ThemeDef(
        "THEME_RULE_OF_LAW_INVERSION",
        "Rule-of-Law Inversion",
        3,
        "Independent enforcement is recast as partisan abuse while allies are shielded and opponents are targeted.",
        "A democratic system depends on law being applied by institutions that are not simply instruments of the ruler. Across the year, prosecutions, pardons, purges, and investigative priorities repeatedly shift in ways that protect regime allies, weaken corruption controls, and redefine accountability as persecution.",
        (r"doj", r"fbi", r"pardons?", r"weaponization", r"foreign influence", r"kleptocracy", r"public integrity", r"corruption", r"law enforcement"),
    ),
    ThemeDef(
        "THEME_CIVIL_SERVICE_PURGE",
        "Civil Service as Target and Terrain",
        4,
        "The professional bureaucracy is weakened, politicized, or hollowed out so it can be replaced by loyalists, contractors, or fear-driven compliance.",
        "The struggle over personnel is not administrative trivia. It determines whether government acts through law, expertise, and continuity, or through personal loyalty, intimidation, and ideological capture.",
        (r"schedule [fg]", r"probationary", r"buyout", r"deferred resignation", r"mass firings?", r"purges?", r"administrative leave", r"reclassification", r"career officials?"),
    ),
    ThemeDef(
        "THEME_INFORMATION_CONTROL",
        "Information and Memory Control",
        5,
        "Data, history, media access, and the public factual baseline are repeatedly manipulated as instruments of power.",
        "This book is not only about policy decisions. It is also about who gets to define reality. When records vanish, media access is reshuffled, public history is rewritten, and disfavored facts are suppressed, democratic self-government loses the common ground it requires.",
        (r"information", r"memory", r"media", r"press", r"history", r"data", r"scrubbed", r"rewrit", r"censorship", r"statistical", r"voice of america", r"public broadcasting"),
    ),
    ThemeDef(
        "THEME_ATTACKS_ON_DISSENT",
        "Attacks on Dissent",
        6,
        "Speech, protest, academic independence, and civil-society opposition are treated as threats to be chilled, punished, or administratively constrained.",
        "A democracy depends on dissent not merely being legal but being politically survivable. Across the book, protest, campus speech, critical media, and civic opposition face escalating pressure through funding threats, prosecutions, immigration tools, and policing.",
        (r"protest crackdown", r"campus protest", r"student visa", r"antisemitism investigation", r"public broadcasting", r"detained student", r"media outlet", r"no kings", r"press credentials?"),
    ),
    ThemeDef(
        "THEME_SELECTIVE_ENFORCEMENT",
        "Selective Enforcement",
        7,
        "State power is applied unevenly depending on loyalty, vulnerability, usefulness, or ideological fit.",
        "One of the clearest democratic warning signs is not simply harsher enforcement, but unequal enforcement. Throughout the year, coercive authority appears repeatedly to fall hardest on the weak, the disfavored, and the politically exposed, while allies and aligned constituencies are treated differently.",
        (r"selective enforcement", r"retaliat", r"punish perceived opponents", r"double standard", r"favored treatment", r"exempt(?:ion|ed) allies", r"make an example"),
    ),
    ThemeDef(
        "THEME_IMMIGRATION_AS_POWER_TEST",
        "Immigration as a Test Case for Power",
        8,
        "Immigration becomes the proving ground for exceptional powers, reduced due process, and coercive governance at scale.",
        "The book shows immigration policy functioning as more than border policy. It becomes the administration’s most elastic field for detention, expulsion, militarization, surveillance, and jurisdictional experimentation, with consequences that spill into the wider constitutional order.",
        (r"ice", r"cbp", r"guant[aá]namo", r"alien enemies act", r"deport", r"detention", r"birthright citizenship", r"asylum", r"expedited removal"),
    ),
    ThemeDef(
        "THEME_TRANSACTIONAL_FOREIGN_POLICY",
        "Transactional Nationalism",
        9,
        "Foreign policy, trade, aid, and security are repeatedly used as unilateral bargaining tools rather than constrained public policy.",
        "International actions in this book often mirror domestic ones: shock, centralization, and leverage. Aid, tariffs, sanctions, territorial rhetoric, and multilateral withdrawals are used not as elements of a stable rules-based order but as executive instruments of pressure and spectacle.",
        (r"tariffs?", r"foreign aid", r"usaid", r"icc", r"multilateral", r"gaza", r"panama", r"greenland", r"sanctions?", r"security assistance"),
    ),
    ThemeDef(
        "THEME_COURTS_AS_COUNTERWEIGHT",
        "Courts as Counterweight",
        10,
        "Judges, litigation, and procedural challenges repeatedly emerge as the principal institutional brakes on executive acceleration.",
        "The judiciary is not portrayed here as a guaranteed savior. But week after week, it appears as one of the few institutions still capable of slowing, clarifying, or temporarily blocking executive moves that other branches will not stop in time.",
        (r"injunction", r"lawsuit", r"litigation", r"temporary block", r"restraining order", r"preliminary injunction", r"appeals court", r"supreme court"),
    ),
    ThemeDef(
        "THEME_LOYALTY_OVER_COMPETENCE",
        "Loyalty Over Competence",
        11,
        "Appointments, restructurings, and removals repeatedly favor ideological obedience over expertise, independence, or administrative capacity.",
        "A recurring pattern in the book is not just that institutions are politicized, but that competence itself is downgraded. This matters because democratic governance depends on capable institutions, not merely obedient ones.",
        (r"confirmed", r"ideologue", r"loyal", r"friendlier outlets", r"industry-aligned", r"project 2025", r"appointed", r"fired inspectors general"),
    ),
    ThemeDef(
        "THEME_NORMALIZATION_THROUGH_REPETITION",
        "Normalization Through Repetition",
        12,
        "Methods that begin as shocks gradually become routine through repetition, fatigue, and the steady lowering of expectations.",
        "Part of the book’s force comes from chronology itself. Similar moves recur across weeks—purges, threats, denials, emergency claims, retaliations—until once-exceptional conduct starts to look ordinary. This appendix gathers those repetitions back into view.",
        (r"continued", r"deepened", r"escalated", r"normalized", r"routine", r"expanded again", r"renewed crackdown", r"another round"),
    ),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the Trump Chronicles Key Themes back-matter chapter.")
    parser.add_argument("--week", type=int, required=True, help="Starting week number.")
    parser.add_argument("--weeks", type=int, default=53, help="Number of consecutive weeks to process.")
    parser.add_argument("--force", action="store_true", help="Overwrite existing outputs.")
    parser.add_argument("--dry-run", action="store_true", help="Show planned work without writing files.")
    parser.add_argument(
        "--step3-root",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Step 3/Weeks"),
        help="Root folder containing Week N subfolders.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/10_BackMatter"),
        help="Directory for key_themes outputs.",
    )
    parser.add_argument("--created-by", type=str, default="build_key_themes.py")
    parser.add_argument("--prompt-version", type=str, default="deterministic-v1")
    parser.add_argument("--model", type=str, default=None)
    parser.add_argument("--git-commit", type=str, default=None)
    parser.add_argument(
        "--log-level",
        "--level",
        dest="log_level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="Logging level.",
    )
    return parser.parse_args()


def configure_logging(level: str) -> None:
    logging.basicConfig(level=getattr(logging, level.upper()), format="%(asctime)s [%(levelname)s] %(message)s")


def week_numbers(start: int, count: int) -> List[int]:
    if start < 1:
        raise ValueError("--week must be >= 1")
    if count < 1:
        raise ValueError("--weeks must be >= 1")
    return list(range(start, start + count))


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        payload = json.load(fh)
    if not isinstance(payload, dict):
        raise ValueError(f"{path}: expected JSON object")
    logging.debug("Loaded %s with keys: %s", path, ", ".join(sorted(payload.keys())))
    return payload


def ensure_writable(output_dir: Path, force: bool) -> Dict[str, Path]:
    json_path = output_dir / "key_themes.json"
    md_path = output_dir / "key_themes.md"
    docx_path = output_dir / "key_themes.docx"
    existing = [p for p in (json_path, md_path, docx_path) if p.exists()]
    if existing and not force:
        raise FileExistsError(f"Output file(s) already exist: {', '.join(str(p) for p in existing)}. Re-run with --force.")
    return {"json": json_path, "md": md_path, "docx": docx_path}


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        text = value
    else:
        text = json.dumps(value, ensure_ascii=False)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def collect_text_fragments(obj: Any) -> List[str]:
    out: List[str] = []
    if obj is None:
        return out
    if isinstance(obj, str):
        cleaned = normalize_text(obj)
        if cleaned:
            out.append(cleaned)
        return out
    if isinstance(obj, list):
        for item in obj:
            out.extend(collect_text_fragments(item))
        return out
    if isinstance(obj, dict):
        for value in obj.values():
            out.extend(collect_text_fragments(value))
        return out
    return out


def analytic_path(root: Path, week: int) -> Path:
    return root / f"Week {week}" / f"weekly_analytic_brief_week{week}.json"


def metadata_path(root: Path, week: int) -> Path:
    return root / f"Week {week}" / f"weekly_digest_metadata_stack_week{week}.json"


def spine_path(root: Path, week: int) -> Path:
    return root / f"Week {week}" / f"weekly_digest_spine_week{week}.json"


def extract_week_title(metadata: Dict[str, Any], spine: Dict[str, Any], week: int) -> str:
    title = (
        metadata.get("title")
        or metadata.get("substack_metadata", {}).get("title")
        or spine.get("article_title")
        or f"Week {week}"
    )
    title = normalize_text(title)
    prefix = "This Week in Democracy:"
    if title.startswith(prefix):
        title = title[len(prefix):].strip()
    return title


def week_source_bundle(step3_root: Path, week: int) -> Dict[str, Any]:
    apath = analytic_path(step3_root, week)
    mpath = metadata_path(step3_root, week)
    spath = spine_path(step3_root, week)

    if not apath.exists():
        raise FileNotFoundError(f"Missing analytic brief for Week {week}: {apath}")
    if not mpath.exists():
        raise FileNotFoundError(f"Missing metadata stack for Week {week}: {mpath}")
    if not spath.exists():
        raise FileNotFoundError(f"Missing weekly spine for Week {week}: {spath}")

    analytic = load_json(apath)
    metadata = load_json(mpath)
    spine = load_json(spath)

    return {
        "week": week,
        "analytic_path": str(apath),
        "metadata_path": str(mpath),
        "spine_path": str(spath),
        "analytic": analytic,
        "metadata": metadata,
        "spine": spine,
        "title": extract_week_title(metadata, spine, week),
    }


def build_search_texts(bundle: Dict[str, Any]) -> Dict[str, str]:
    analytic = bundle["analytic"]
    metadata = bundle["metadata"]
    spine = bundle["spine"]

    analytic_parts: List[str] = []
    analytic_parts.extend(collect_text_fragments(analytic.get("summary")))
    analytic_parts.extend(collect_text_fragments(analytic.get("top_themes")))
    analytic_parts.extend(collect_text_fragments(analytic.get("institutional_shifts")))

    metadata_parts: List[str] = []
    metadata_parts.extend(collect_text_fragments(metadata.get("title")))
    metadata_parts.extend(collect_text_fragments(metadata.get("subtitle")))
    metadata_parts.extend(collect_text_fragments(metadata.get("week_summary")))
    metadata_parts.extend(collect_text_fragments(metadata.get("category_flags")))
    metadata_parts.extend(collect_text_fragments(metadata.get("internal_tags")))
    metadata_parts.extend(collect_text_fragments(metadata.get("whole_week_arcs")))

    spine_parts: List[str] = []
    spine_parts.extend(collect_text_fragments(spine.get("article_title")))
    spine_parts.extend(collect_text_fragments(spine.get("article_subtitle")))
    spine_parts.extend(collect_text_fragments(spine.get("sections")))

    return {
        "analytic": " ".join(analytic_parts),
        "metadata": " ".join(metadata_parts),
        "spine": " ".join(spine_parts),
    }

def count_matched_forms(theme: ThemeDef, search_texts: Dict[str, str]) -> int:
    form_groups = THEME_FORM_GROUPS.get(theme.theme_id)
    if not form_groups:
        return 0

    matched_forms = 0
    for _form_name, patterns in form_groups:
        form_matched = False
        for pattern in patterns:
            regex = re.compile(pattern, flags=re.IGNORECASE)
            if any(text and regex.search(text) for text in search_texts.values()):
                form_matched = True
                break
        if form_matched:
            matched_forms += 1
    return matched_forms

def score_theme(theme: ThemeDef, search_texts: Dict[str, str]) -> Tuple[int, List[str], int]:
    total = 0
    reasons: List[str] = []
    matched_sources: set[str] = set()
    weights = {"analytic": 4, "metadata": 2, "spine": 3}

    for pattern in theme.patterns:
        regex = re.compile(pattern, flags=re.IGNORECASE)
        matched_any = False
        for source_name, text in search_texts.items():
            if text and regex.search(text):
                total += weights[source_name]
                matched_sources.add(source_name)
                matched_any = True
        if matched_any:
            reasons.append(pattern)

    return total, reasons, len(matched_sources)



# Helper for theme-specific rules
def get_theme_rules(theme_id: str) -> Dict[str, int]:
    return {
        "primary_min_score": PRIMARY_MIN_SCORE,
        "primary_min_patterns": PRIMARY_MIN_PATTERNS,
        "primary_min_sources": PRIMARY_MIN_SOURCES,
        "secondary_min_score": SECONDARY_MIN_SCORE,
        "secondary_min_patterns": SECONDARY_MIN_PATTERNS,
        "supporting_min_score": SUPPORTING_MIN_SCORE,
        "supporting_min_patterns": SUPPORTING_MIN_PATTERNS,
        "target_primary_weeks": TARGET_PRIMARY_WEEKS,
        "target_total_weeks": TARGET_TOTAL_WEEKS,
        "fallback_total_weeks": FALLBACK_TOTAL_WEEKS,
        **THEME_RULE_OVERRIDES.get(theme_id, {}),
        "primary_min_forms": PRIMARY_MIN_FORMS,
        "secondary_min_forms": SECONDARY_MIN_FORMS,
        "supporting_min_forms": SUPPORTING_MIN_FORMS,
    }

def relevance_from_score(theme_id: str, score: int, pattern_count: int, source_count: int, form_count: int) -> Optional[str]:
    rules = get_theme_rules(theme_id)
    if (
        score >= rules["primary_min_score"]
        and pattern_count >= rules["primary_min_patterns"]
        and source_count >= rules["primary_min_sources"]
        and form_count >= rules["primary_min_forms"]
    ):
        return "primary"
    if (
        score >= rules["secondary_min_score"]
        and pattern_count >= rules["secondary_min_patterns"]
        and form_count >= rules["secondary_min_forms"]
    ):
        return "secondary"
    if (
        score >= rules["supporting_min_score"]
        and pattern_count >= rules["supporting_min_patterns"]
        and form_count >= rules["supporting_min_forms"]
    ):
        return "supporting"
    return None


# Helper functions for theme locator discipline
def source_name_matches(patterns: Sequence[str], text: str) -> bool:
    return any(re.search(pattern, text, flags=re.IGNORECASE) for pattern in patterns)


def select_locator_weeks(theme_id: str, locator_candidates: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not locator_candidates:
        return []

    rules = get_theme_rules(theme_id)
    ranked = sorted(
        locator_candidates,
        key=lambda item: (
            RELEVANCE_RANK[item["relevance"]],
            -item["form_count"],
            -item["score"],
            -item["pattern_count"],
            -item["source_count"],
            item["week"],
        ),
    )

    primaries = [item for item in ranked if item["relevance"] == "primary"]
    secondaries = [item for item in ranked if item["relevance"] == "secondary"]
    supportings = [item for item in ranked if item["relevance"] == "supporting"]

    selected: List[Dict[str, Any]] = primaries[: rules["target_primary_weeks"]]

    if len(selected) < rules["target_primary_weeks"]:
        needed = rules["target_total_weeks"] - len(selected)
        if needed > 0:
            selected.extend(secondaries[:needed])

    if not selected:
        selected = secondaries[: rules["fallback_total_weeks"]]

    if not selected:
        selected = supportings[: rules["fallback_total_weeks"]]

    selected = sorted(selected, key=lambda item: item["week"])
    return [{"week": item["week"], "relevance": item["relevance"]} for item in selected]


def build_theme_hits(bundles: Sequence[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    theme_hits: Dict[str, Dict[str, Any]] = {
        theme.theme_id: {
            "definition": theme,
            "locator_weeks": [],
            "analytic_brief_weeks": [],
            "metadata_stack_weeks": [],
            "weekly_spine_weeks": [],
            "notes": [],
        }
        for theme in THEME_CATALOGUE
    }

    for bundle in bundles:
        week = bundle["week"]
        texts = build_search_texts(bundle)

        for theme in THEME_CATALOGUE:
            score, reasons, source_count = score_theme(theme, texts)
            pattern_count = len(reasons)
            form_count = count_matched_forms(theme, texts)
            relevance = relevance_from_score(theme.theme_id, score, pattern_count, source_count, form_count)
            if relevance is None:
                continue

            theme_hits[theme.theme_id]["locator_weeks"].append(
                {
                    "week": week,
                    "relevance": relevance,
                    "score": score,
                    "pattern_count": pattern_count,
                    "source_count": source_count,
                    "form_count": form_count,
                }
            )

            if source_name_matches(theme.patterns, texts["analytic"]):
                theme_hits[theme.theme_id]["analytic_brief_weeks"].append(week)
            if source_name_matches(theme.patterns, texts["metadata"]):
                theme_hits[theme.theme_id]["metadata_stack_weeks"].append(week)
            if source_name_matches(theme.patterns, texts["spine"]):
                theme_hits[theme.theme_id]["weekly_spine_weeks"].append(week)

            if reasons:
                theme_hits[theme.theme_id]["notes"].append(
                    f"Week {week}: matched {', '.join(reasons[:6])}"
                )

            logging.debug(
                "Week %s matched %s with score=%s pattern_count=%s source_count=%s form_count=%s relevance=%s",
                week,
                theme.theme_name,
                score,
                pattern_count,
                source_count,
                form_count,
                relevance,
            )

    return theme_hits


def collapse_week_ranges(weeks: Sequence[int]) -> str:
    if not weeks:
        return ""

    ordered = sorted(set(weeks))
    ranges: List[str] = []
    start = prev = ordered[0]

    for week in ordered[1:]:
        if week == prev + 1:
            prev = week
            continue
        ranges.append(f"Week {start}" if start == prev else f"Weeks {start}–{prev}")
        start = prev = week

    ranges.append(f"Week {start}" if start == prev else f"Weeks {start}–{prev}")
    return ", ".join(ranges)


def build_theme_entries(theme_hits: Dict[str, Dict[str, Any]]) -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []

    for theme in sorted(THEME_CATALOGUE, key=lambda t: t.sort_order):
        hit = theme_hits[theme.theme_id]
        if not hit["locator_weeks"]:
            continue

        locator_weeks = select_locator_weeks(theme.theme_id, hit["locator_weeks"])
        if not locator_weeks:
            continue
        entries.append(
            {
                "theme_id": theme.theme_id,
                "theme_name": theme.theme_name,
                "sort_order": theme.sort_order,
                "framing_line": theme.framing_line,
                "body": theme.body,
                "locator_weeks": locator_weeks,
                "source_summary": {
                    "analytic_brief_weeks": sorted(set(hit["analytic_brief_weeks"])),
                    "metadata_stack_weeks": sorted(set(hit["metadata_stack_weeks"])),
                    "weekly_spine_weeks": sorted(set(hit["weekly_spine_weeks"])),
                },
                "notes": hit["notes"][:12],
            }
        )

    return entries


def build_output(
    entries: List[Dict[str, Any]],
    step3_root: Path,
    created_by: str,
    prompt_version: str,
    model: Optional[str],
    git_commit: Optional[str],
    timeline_file: Optional[str] = None,
) -> Dict[str, Any]:
    now_utc = datetime.now(UTC)
    run_id = f"key-themes-{now_utc.strftime('%Y%m%dT%H%M%SZ')}-{uuid.uuid4().hex[:8]}"

    return {
        "schema_name": SCHEMA_NAME,
        "schema_version": SCHEMA_VERSION,
        "section": SECTION_NAME,
        "book_title": BOOK_TITLE,
        "build": {
            "created_at": now_utc.replace(microsecond=0).isoformat().replace("+00:00", "Z"),
            "created_by": created_by,
            "model": model,
            "run_id": run_id,
            "prompt_version": prompt_version,
            "git_commit": git_commit,
        },
        "generated_from": {
            "analytic_brief_pattern": str(step3_root / "Week N" / "weekly_analytic_brief_weekN.json"),
            "metadata_stack_pattern": str(step3_root / "Week N" / "weekly_digest_metadata_stack_weekN.json"),
            "weekly_spine_pattern": str(step3_root / "Week N" / "weekly_digest_spine_weekN.json"),
            "timeline_by_week_file": timeline_file,
        },
        "themes": entries,
    }


def render_markdown(payload: Dict[str, Any]) -> str:
    lines: List[str] = ["# Key Themes", ""]

    for theme in payload["themes"]:
        lines.append(f"## {theme['theme_name']}")
        lines.append("")
        lines.append(theme["framing_line"])
        lines.append("")
        lines.append(theme["body"])
        lines.append("")
        locator_text = collapse_week_ranges([loc["week"] for loc in theme["locator_weeks"]])
        lines.append(f"**Prominent in:** {locator_text}")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def render_docx(markdown_text: str, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()

    for block in markdown_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        if block.startswith("# "):
            p = doc.add_paragraph()
            p.style = "Heading 1"
            p.add_run(block[2:].strip())
            continue

        if block.startswith("## "):
            p = doc.add_paragraph()
            p.style = "Heading 2"
            p.add_run(block[3:].strip())
            continue

        if block.startswith("**Prominent in:**"):
            p = doc.add_paragraph()
            run1 = p.add_run("Prominent in: ")
            run1.bold = True
            p.add_run(block[len("**Prominent in:**"):].strip())
            continue

        doc.add_paragraph(block)

    doc.save(str(output_path))


def main() -> int:
    args = parse_args()
    configure_logging(args.log_level)

    weeks = week_numbers(args.week, args.weeks)
    outputs = {
        "json": args.output_dir / "key_themes.json",
        "md": args.output_dir / "key_themes.md",
        "docx": args.output_dir / "key_themes.docx",
    }

    logging.info("Building Key Themes for Weeks %s to %s", weeks[0], weeks[-1])
    logging.info("Step 3 root: %s", args.step3_root)
    logging.info("Output dir: %s", args.output_dir)

    if args.dry_run:
        for week in weeks:
            logging.info("[DRY RUN] Would read: %s", analytic_path(args.step3_root, week))
            logging.info("[DRY RUN] Would read: %s", metadata_path(args.step3_root, week))
            logging.info("[DRY RUN] Would read: %s", spine_path(args.step3_root, week))
        for path in outputs.values():
            logging.info("[DRY RUN] Would write: %s", path)
        return 0

    try:
        args.output_dir.mkdir(parents=True, exist_ok=True)
        ensure_writable(args.output_dir, args.force)
    except Exception as exc:
        logging.error(str(exc))
        return 1

    bundles: List[Dict[str, Any]] = []
    failed_weeks: List[int] = []

    for week in weeks:
        try:
            bundles.append(week_source_bundle(args.step3_root, week))
            logging.info("Loaded Week %s", week)
        except Exception as exc:
            logging.error("Failed Week %s: %s", week, exc)
            failed_weeks.append(week)

    if not bundles:
        logging.error("No valid weeks loaded; nothing to write.")
        return 1

    theme_hits = build_theme_hits(bundles)
    entries = build_theme_entries(theme_hits)

    if not entries:
        logging.error("No themes identified; nothing to write.")
        return 1

    payload = build_output(
        entries=entries,
        step3_root=args.step3_root,
        created_by=args.created_by,
        prompt_version=args.prompt_version,
        model=args.model,
        git_commit=args.git_commit,
    )
    markdown = render_markdown(payload)

    try:
        with outputs["json"].open("w", encoding="utf-8") as fh:
            json.dump(payload, fh, indent=2, ensure_ascii=False)
            fh.write("\n")

        with outputs["md"].open("w", encoding="utf-8") as fh:
            fh.write(markdown)

        render_docx(markdown, outputs["docx"])
    except Exception as exc:
        logging.error("Failed writing outputs: %s", exc)
        return 1

    logging.info("WROTE: %s", outputs["json"])
    logging.info("WROTE: %s", outputs["md"])
    logging.info("WROTE: %s", outputs["docx"])

    if failed_weeks:
        logging.warning("Skipped failed weeks: %s", ", ".join(map(str, failed_weeks)))

    return 0


if __name__ == "__main__":
    sys.exit(main())