#!/usr/bin/env python3
"""
build_timeline_by_week.py

Metadata-only first pass for the Trump Chronicles back-matter chapter
"Timeline by Week".

Behavior:
- requires --week N
- optional --weeks K means K consecutive weeks starting from --week
- optional --force overwrites existing output files
- without --force, existing final outputs cause the run to stop before writing
- reads only weekly_digest_metadata_stack_weekNN.json
- uses metadata week_summary verbatim as gloss
- writes one aggregated JSON, one markdown chapter, one DOCX chapter
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore


SCHEMA_NAME = "phase2_timeline_output"
SCHEMA_VERSION = "1.0"
SECTION_NAME = "timeline_by_week"
BOOK_TITLE = "Trump Chronicles"
TITLE_PREFIX = "This Week in Democracy:"


@dataclass(frozen=True)
class WeekPaths:
    week_number: int
    metadata_path: Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build the Trump Chronicles Timeline by Week back-matter chapter."
    )
    parser.add_argument("--week", type=int, required=True, help="Starting week number.")
    parser.add_argument(
        "--weeks",
        type=int,
        default=1,
        help="Number of consecutive weeks to process starting from --week.",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite existing output files.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show planned work without writing files.",
    )
    parser.add_argument(
        "--step3-root",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Step 3/Weeks"),
        help="Root folder containing Week N subfolders.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/10_BackMatter"),
        help="Directory for timeline_by_week outputs.",
    )
    parser.add_argument(
        "--created-by",
        type=str,
        default="build_timeline_by_week.py",
        help="Value for build.created_by in the output JSON.",
    )
    parser.add_argument(
        "--prompt-version",
        type=str,
        default="metadata-only-v1",
        help="Value for build.prompt_version in the output JSON.",
    )
    parser.add_argument(
        "--model",
        type=str,
        default=None,
        help="Value for build.model in the output JSON. Use null for metadata-only runs.",
    )
    parser.add_argument(
        "--git-commit",
        type=str,
        default=None,
        help="Optional git commit hash for provenance.",
    )
    parser.add_argument(
        "--log-level",
        "--level",
        dest="log_level",
        type=str,
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        help="Logging level.",
    )
    return parser.parse_args()


def configure_logging(level: str) -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper()),
        format="%(asctime)s [%(levelname)s] %(message)s",
    )


def week_range(start_week: int, count: int) -> List[int]:
    if start_week < 1:
        raise ValueError("--week must be >= 1")
    if count < 1:
        raise ValueError("--weeks must be >= 1")
    return list(range(start_week, start_week + count))


def metadata_path_for_week(step3_root: Path, week_number: int) -> Path:
    return step3_root / f"Week {week_number}" / f"weekly_digest_metadata_stack_week{week_number}.json"


def collect_week_paths(step3_root: Path, start_week: int, count: int) -> List[WeekPaths]:
    paths: List[WeekPaths] = []
    for week_number in week_range(start_week, count):
        metadata_path = metadata_path_for_week(step3_root, week_number)
        logging.debug("Planned Week %s metadata path: %s", week_number, metadata_path)
        paths.append(WeekPaths(week_number=week_number, metadata_path=metadata_path))
    return paths




# New helper: load_existing_entries
def load_existing_entries(json_path: Path) -> List[Dict[str, Any]]:
    if not json_path.exists():
        logging.debug("No existing timeline JSON found at %s", json_path)
        return []

    payload = load_json(json_path)
    if not isinstance(payload, dict):
        raise ValueError(f"{json_path}: existing timeline output is not a JSON object")

    entries = payload.get("entries")
    if entries is None:
        logging.debug("Existing timeline JSON at %s has no entries array; treating as empty", json_path)
        return []
    if not isinstance(entries, list):
        raise ValueError(f"{json_path}: existing timeline output has invalid 'entries' field")

    valid_entries: List[Dict[str, Any]] = []
    for entry in entries:
        if isinstance(entry, dict) and isinstance(entry.get("week"), int):
            valid_entries.append(entry)
        else:
            logging.warning("Skipping invalid existing entry in %s: %r", json_path, entry)

    logging.debug("Loaded %s existing timeline entrie(s) from %s", len(valid_entries), json_path)
    return valid_entries


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    if isinstance(payload, dict):
        logging.debug("Loaded JSON object from %s with top-level keys: %s", path, ", ".join(sorted(payload.keys())))
    else:
        logging.debug("Loaded JSON from %s with top-level type: %s", path, type(payload).__name__)
    return payload


def require_string(value: Any, field_name: str, source_path: Path) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{source_path}: required non-empty string field '{field_name}' is missing or invalid")
    return value.strip()


def require_dict(value: Any, field_name: str, source_path: Path) -> Dict[str, Any]:
    if not isinstance(value, dict):
        raise ValueError(f"{source_path}: required object field '{field_name}' is missing or invalid")
    return value


def clean_title(raw_title: str) -> str:
    title = raw_title.strip()
    if title.startswith(TITLE_PREFIX):
        title = title[len(TITLE_PREFIX):].strip()
    return title


def parse_iso_date(date_text: str, source_path: Path, field_name: str) -> datetime:
    try:
        return datetime.strptime(date_text, "%Y-%m-%d")
    except ValueError as exc:
        raise ValueError(f"{source_path}: invalid {field_name} '{date_text}' (expected YYYY-MM-DD)") from exc


def format_us_date_range(start_date: str, end_date: str, source_path: Path) -> str:
    start_dt = parse_iso_date(start_date, source_path, "window.start_date")
    end_dt = parse_iso_date(end_date, source_path, "window.end_date")

    if start_dt.year != end_dt.year:
        return f"{start_dt.strftime('%B')} {start_dt.day}, {start_dt.year}\u2013{end_dt.strftime('%B')} {end_dt.day}, {end_dt.year}"

    if start_dt.month == end_dt.month:
        return f"{start_dt.strftime('%B')} {start_dt.day}\u2013{end_dt.day}, {start_dt.year}"

    return f"{start_dt.strftime('%B')} {start_dt.day}\u2013{end_dt.strftime('%B')} {end_dt.day}, {start_dt.year}"


def normalize_gloss(text: str) -> str:
    # First pass: no compression, but normalize whitespace and preserve paragraph intent.
    # Double newlines become spaces so the timeline remains one compact paragraph per entry.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n\s*\n+", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def build_entry(week_number: int, metadata_path: Path) -> Dict[str, Any]:
    metadata = load_json(metadata_path)
    logging.debug("Building timeline entry from %s", metadata_path)

    raw_title = require_string(metadata.get("title"), "title", metadata_path)
    week_summary = require_string(metadata.get("week_summary"), "week_summary", metadata_path)

    window = require_dict(metadata.get("window"), "window", metadata_path)
    start_date = require_string(window.get("start_date"), "window.start_date", metadata_path)
    end_date = require_string(window.get("end_date"), "window.end_date", metadata_path)

    chapter_title = clean_title(raw_title)
    date_display = format_us_date_range(start_date, end_date, metadata_path)
    gloss = normalize_gloss(week_summary)

    logging.debug(
        "Week %s entry fields: title=%r, start_date=%s, end_date=%s, gloss_chars=%s",
        week_number,
        chapter_title,
        start_date,
        end_date,
        len(gloss),
    )

    return {
        "week": week_number,
        "start_date": start_date,
        "end_date": end_date,
        "date_display": date_display,
        "chapter_title": chapter_title,
        "gloss": gloss,
        "source_priority_used": "metadata.week_summary",
        "source_metadata_stack": str(metadata_path),
        "source_weekly_spine": None,
    }


def build_output(
    entries: List[Dict[str, Any]],
    step3_root: Path,
    created_by: str,
    prompt_version: str,
    model: Optional[str],
    git_commit: Optional[str],
) -> Dict[str, Any]:
    now_utc = datetime.now(UTC)
    run_id = f"timeline-{now_utc.strftime('%Y%m%dT%H%M%SZ')}-{uuid.uuid4().hex[:8]}"
    return {
        "schema_name": SCHEMA_NAME,
        "schema_version": SCHEMA_VERSION,
        "section": SECTION_NAME,
        "book_title": BOOK_TITLE,
        "build": {
            "created_at": now_utc.replace(microsecond=0).isoformat().replace("+00:00", "Z"),
            "created_by": created_by,
            "model": model,
            "run_id": run_id,
            "prompt_version": prompt_version,
            "git_commit": git_commit,
        },
        "generated_from": {
            "metadata_stack_pattern": str(step3_root / "Week N" / "weekly_digest_metadata_stack_weekN.json"),
            "weekly_spine_pattern": str(step3_root / "Week N" / "weekly_digest_spine_weekN.json"),
        },
        "entries": sorted(entries, key=lambda x: x["week"]),
    }


def render_markdown(payload: Dict[str, Any]) -> str:
    lines: List[str] = []
    lines.append("# Timeline by Week")
    lines.append("")
    for entry in payload["entries"]:
        logging.debug("Rendering markdown entry for Week %s", entry["week"])
        lines.append(
            f"**Week {entry['week']}: {entry['date_display']} \u2014 {entry['chapter_title']}**"
        )
        lines.append(entry["gloss"])
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def render_docx(markdown_text: str, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed in this environment. Install it to write the DOCX output."
        )

    doc = Document()

    paragraphs = markdown_text.split("\n\n")
    first_heading_done = False

    for block in paragraphs:
        block = block.strip()
        logging.debug("DOCX block: %r", block[:120])
        if not block:
            continue

        if block.startswith("# "):
            text = block[2:].strip()
            p = doc.add_paragraph()
            p.style = "Heading 1"
            p.add_run(text)
            first_heading_done = True
            continue

        if block.startswith("**") and block.endswith("**"):
            text = block.strip("*")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            continue

        # Handle the normal prose paragraphs.
        doc.add_paragraph(block)

    if not first_heading_done:
        p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        p.style = "Heading 1"

    doc.save(str(output_path))


def main() -> int:
    args = parse_args()
    configure_logging(args.log_level)

    week_paths = collect_week_paths(args.step3_root, args.week, args.weeks)
    logging.debug("Resolved %s planned week path(s)", len(week_paths))
    output_paths = {
        "json": args.output_dir / "timeline_by_week.json",
        "md": args.output_dir / "timeline_by_week.md",
        "docx": args.output_dir / "timeline_by_week.docx",
    }

    logging.info("Building Timeline by Week for Weeks %s to %s", args.week, args.week + args.weeks - 1)
    logging.info("Step 3 root: %s", args.step3_root)
    logging.info("Output dir: %s", args.output_dir)

    if args.dry_run:
        for wp in week_paths:
            logging.info("[DRY RUN] Would read: %s", wp.metadata_path)
        logging.info("[DRY RUN] Would write: %s", output_paths["json"])
        logging.info("[DRY RUN] Would write: %s", output_paths["md"])
        logging.info("[DRY RUN] Would write: %s", output_paths["docx"])
        return 0

    logging.debug("Output targets: json=%s md=%s docx=%s", output_paths["json"], output_paths["md"], output_paths["docx"])

    try:
        args.output_dir.mkdir(parents=True, exist_ok=True)
        existing_entries = load_existing_entries(output_paths["json"])
    except Exception as exc:
        logging.error(str(exc))
        return 1

    existing_entries_by_week: Dict[int, Dict[str, Any]] = {
        entry["week"]: entry for entry in existing_entries
    }
    existing_weeks = set(existing_entries_by_week.keys())
    logging.debug("Existing timeline weeks: %s", sorted(existing_weeks))

    entries: List[Dict[str, Any]] = []
    skipped_existing_weeks: List[int] = []
    replaced_weeks: List[int] = []
    missing_weeks: List[int] = []
    failed_weeks: List[int] = []

    for wp in week_paths:
        logging.debug("Checking Week %s metadata file: %s", wp.week_number, wp.metadata_path)
        if not args.force and wp.week_number in existing_weeks:
            logging.info("Skipping existing Week %s (already present in %s)", wp.week_number, output_paths["json"])
            skipped_existing_weeks.append(wp.week_number)
            continue

        if not wp.metadata_path.exists():
            logging.warning("Missing metadata for Week %s: %s", wp.week_number, wp.metadata_path)
            missing_weeks.append(wp.week_number)
            continue

        try:
            entry = build_entry(wp.week_number, wp.metadata_path)
            entries.append(entry)
            if args.force and wp.week_number in existing_weeks:
                replaced_weeks.append(wp.week_number)
                logging.info("Reprocessed Week %s (--force; replacing existing entry)", wp.week_number)
            else:
                logging.info("Processed Week %s", wp.week_number)
        except Exception as exc:
            logging.error("Failed Week %s: %s", wp.week_number, exc)
            failed_weeks.append(wp.week_number)

    merged_entries_by_week: Dict[int, Dict[str, Any]] = dict(existing_entries_by_week)
    for entry in entries:
        merged_entries_by_week[entry["week"]] = entry

    if not merged_entries_by_week:
        logging.error("No valid timeline entries exist; nothing to write.")
        return 1

    if not entries:
        logging.info("No new weeks required processing; existing timeline outputs will be preserved and refreshed.")

    payload = build_output(
        entries=list(merged_entries_by_week.values()),
        step3_root=args.step3_root,
        created_by=args.created_by,
        prompt_version=args.prompt_version,
        model=args.model,
        git_commit=args.git_commit,
    )
    logging.debug("Built payload with %s entry/entries", len(payload["entries"]))

    markdown_text = render_markdown(payload)
    logging.debug("Rendered markdown length: %s characters", len(markdown_text))

    try:
        logging.debug("Writing output files")
        with output_paths["json"].open("w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
            f.write("\n")

        with output_paths["md"].open("w", encoding="utf-8") as f:
            f.write(markdown_text)

        render_docx(markdown_text, output_paths["docx"])
    except Exception as exc:
        logging.error("Failed to write outputs: %s", exc)
        return 1

    logging.info("WROTE: %s", output_paths["json"])
    logging.info("WROTE: %s", output_paths["md"])
    logging.info("WROTE: %s", output_paths["docx"])

    if missing_weeks:
        logging.warning("Skipped missing weeks: %s", ", ".join(map(str, missing_weeks)))
    if failed_weeks:
        logging.warning("Skipped failed weeks: %s", ", ".join(map(str, failed_weeks)))

    if skipped_existing_weeks:
        logging.info("Skipped already-built weeks: %s", ", ".join(map(str, skipped_existing_weeks)))
    if replaced_weeks:
        logging.info("Replaced existing weeks due to --force: %s", ", ".join(map(str, replaced_weeks)))

    return 0


if __name__ == "__main__":
    sys.exit(main())