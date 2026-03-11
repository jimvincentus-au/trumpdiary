#!/usr/bin/env python3
"""
build_vellum_import_documents.py

Generate Vellum import artifacts from:
- source chapters in: /Trump Diary/08_RewrittenChapters/
- metadata in:       /Step 3/Weeks/Week N/metadata_stack_weekN.json

Outputs:
- enriched markdown with YAML front matter
- .docx files for Vellum import

Design decisions:
- source files are never modified
- title comes from metadata["title"]
- subtitle comes from metadata["subtitle"]
- subtitle is rendered as an opening block quotation, not a Vellum subtitle element
- filenames look like: week001_slug-of-up-to-five-words.md/.docx
- slug is derived from title after removing constant boilerplate
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from datetime import date
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional, cast

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle

logger = logging.getLogger(__name__)


# -----------------------------
# Configuration
# -----------------------------

TITLE_PREFIXES_TO_STRIP = [
    r"^This Week in Democracy:\s*",
    r"^This Week in Democracy\s*[-–—:]\s*",
]

DEFAULT_CHAPTER_GLOB = "week_*_rewritten.md"
DEFAULT_METADATA_NAME_TEMPLATE = "weekly_digest_metadata_stack_week{week}.json"

STOPWORDS = {
    "a", "an", "and", "as", "at", "by", "for", "from", "in", "into", "of",
    "on", "or", "the", "to", "with", "without", "over", "under", "after",
    "before", "during", "week", "democracy", "clock", "this", "that",
}

FILTER_STOPWORDS_IN_SLUG = True


# -----------------------------
# Data model
# -----------------------------

class OverwriteRequiredError(Exception):
    """Raised when output files already exist and --force was not supplied."""


@dataclass
class ChapterJob:
    week_num: int
    source_md: Path
    source_metadata: Path
    title: str
    visible_title: str
    subtitle: str
    period_label: str
    start_date: str
    end_date: str
    slug: str
    week_display: str


# -----------------------------
# Helpers
# -----------------------------

def extract_week_num_from_filename(path: Path) -> int:
    """
    Extract week number from filenames like:
    - week_01_rewritten.md
    - week_1_rewritten.md
    """
    m = re.search(r"week_(\d+)_rewritten\.md$", path.name, flags=re.IGNORECASE)
    if not m:
        raise ValueError(f"Could not extract week number from filename: {path.name}")
    return int(m.group(1))


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def strip_constant_title_prefix(title: str) -> str:
    result = title.strip()
    for pattern in TITLE_PREFIXES_TO_STRIP:
        result = re.sub(pattern, "", result, flags=re.IGNORECASE)
    return result.strip()


# Helper to derive the visible chapter title without boilerplate prefix
def visible_title_from_metadata(title: str) -> str:
    visible = strip_constant_title_prefix(title)
    return visible or title.strip()


# Helper for week/period display line
def week_display_from_metadata(
    week_num: int,
    period_label: str,
    start_date: str,
    end_date: str,
) -> str:
    week_label = f"Week {week_num}"

    if start_date and end_date:
        try:
            start = date.fromisoformat(start_date)
            end = date.fromisoformat(end_date)

            if start.year == end.year:
                if start.month == end.month:
                    date_range = f"{start.strftime('%B')} {start.day}\u2013{end.day}, {start.year}"
                else:
                    date_range = (
                        f"{start.strftime('%B')} {start.day}\u2013"
                        f"{end.strftime('%B')} {end.day}, {start.year}"
                    )
            else:
                date_range = (
                    f"{start.strftime('%B')} {start.day}, {start.year}\u2013"
                    f"{end.strftime('%B')} {end.day}, {end.year}"
                )

            return f"{week_label}: {date_range}"
        except ValueError:
            pass

    cleaned_period = period_label.strip()
    if not cleaned_period:
        return week_label
    if cleaned_period.lower().startswith(week_label.lower()):
        return cleaned_period
    return f"{week_label}: {cleaned_period}"


def normalize_quotes_dashes(text: str) -> str:
    """
    Mild normalization for filenames/metadata handling.
    Leaves content readable.
    """
    if not text:
        return text
    replacements = {
        "\u2018": "",
        "\u2019": "",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00A0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def slugify_phrase(text: str, max_words: int = 5) -> str:
    """
    Slugify after removing boilerplate.
    Example:
    "This Week in Democracy: Emergencies as Governing Method"
    -> "emergencies-as-governing-method"
    """
    text = strip_constant_title_prefix(normalize_quotes_dashes(text))
    text = text.lower()

    text = re.sub(r"[^a-z0-9\s-]", " ", text)
    words = [w.strip(" -") for w in text.split() if w.strip(" -")]

    if FILTER_STOPWORDS_IN_SLUG:
        filtered = [w for w in words if w not in STOPWORDS]
        if filtered:
            words = filtered

    words = words[:max_words]
    slug = "-".join(words)
    slug = re.sub(r"-{2,}", "-", slug).strip("-")

    return slug or "chapter"


def yaml_escape(value: str) -> str:
    value = value.replace("\\", "\\\\").replace('"', '\\"')
    return f'"{value}"'


def build_yaml_front_matter(job: ChapterJob) -> str:
    visible_title = job.visible_title
    lines = [
        "---",
        f"title: {yaml_escape(visible_title)}",
        f"week: {job.week_num}",
        f"period_label: {yaml_escape(job.period_label)}",
        f"start_date: {yaml_escape(job.start_date)}",
        f"end_date: {yaml_escape(job.end_date)}",
        f"source_chapter: {yaml_escape(job.source_md.name)}",
        f"source_metadata: {yaml_escape(job.source_metadata.name)}",
        "---",
        "",
    ]
    return "\n".join(lines)


def strip_leading_yaml_and_h1(markdown_text: str, expected_title: Optional[str] = None) -> str:
    """
    Safely remove:
    - existing YAML front matter
    - a leading H1 that would duplicate the injected title
    """
    text = markdown_text.lstrip("\ufeff")

    if text.startswith("---\n") or text.startswith("---\r\n"):
        m = re.match(r"^---\r?\n.*?\r?\n---\r?\n*", text, flags=re.DOTALL)
        if m:
            text = text[m.end():]

    lines = text.splitlines()
    if lines:
        first = lines[0].strip()
        if first.startswith("# "):
            h1 = first[2:].strip()
            if expected_title is None or h1 == expected_title:
                lines = lines[1:]
                if lines and not lines[0].strip():
                    lines = lines[1:]
                text = "\n".join(lines)

    return text.lstrip()


def read_source_body(source_md_path: Path, expected_title: Optional[str] = None) -> str:
    logger.debug("Reading source markdown from %s", source_md_path)
    raw = source_md_path.read_text(encoding="utf-8")
    cleaned = strip_leading_yaml_and_h1(raw, expected_title=expected_title).rstrip() + "\n"
    logger.debug("Prepared source body from %s (%d characters)", source_md_path.name, len(cleaned))
    return cleaned


def markdown_for_output(job: ChapterJob, source_body: str) -> str:
    front_matter = build_yaml_front_matter(job)
    visible = [
        f"# {job.visible_title}",
        "",
        job.week_display,
        "",
        f"> {job.subtitle}",
        "",
        source_body.rstrip(),
        "",
    ]
    return front_matter + "\n".join(visible)


def apply_paragraph_border_left(paragraph) -> None:
    """
    Optional light visual cue in Word for the opening quotation.
    Vellum may or may not preserve it, but the docx remains readable.
    """
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "6")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "808080")
    pbdr.append(left)
    p_pr.append(pbdr)


def set_document_margins(doc: DocxDocumentType) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)



def add_title(doc: DocxDocumentType, title: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True


# Add a DOCX helper for the week/period line, above the block quote
def add_week_display(doc: DocxDocumentType, week_display: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(week_display)


def add_opening_blockquote(doc: DocxDocumentType, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(subtitle)
    run.italic = True
    apply_paragraph_border_left(p)


def add_markdown_body_minimal(doc: DocxDocumentType, body: str) -> None:
    """
    Minimal markdown handling:
    - blank-line separated paragraphs
    - blockquotes
    - ## / ### headings if they occur
    - bullet list items starting with '- ' or '* '
    - everything else as ordinary prose

    This is intentionally conservative for Vellum-friendly import.
    """
    blocks = re.split(r"\n\s*\n", body.strip(), flags=re.MULTILINE)

    for block in blocks:
        lines = block.splitlines()
        if not lines:
            continue

        if len(lines) == 1 and lines[0].startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(lines[0][3:].strip())
            continue

        if len(lines) == 1 and lines[0].startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(lines[0][4:].strip())
            continue

        if all(line.lstrip().startswith(">") for line in lines):
            quote_text = "\n".join(re.sub(r"^\s*>\s?", "", line) for line in lines).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.15)
            run = p.add_run(quote_text)
            run.italic = True
            apply_paragraph_border_left(p)
            continue

        if all(re.match(r"^\s*[-*]\s+", line) for line in lines):
            for line in lines:
                item = re.sub(r"^\s*[-*]\s+", "", line).strip()
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item)
            continue

        if len(lines) == 1 and lines[0].startswith("# "):
            continue

        para_text = "\n".join(line.strip() for line in lines).strip()
        if para_text:
            p = doc.add_paragraph(style="Normal")
            p.add_run(para_text)


def get_paragraph_style(doc: DocxDocumentType, style_name: str) -> ParagraphStyle:
    return cast(ParagraphStyle, doc.styles[style_name])


def write_docx(job: ChapterJob, source_body: str, output_path: Path) -> None:
    doc = Document()
    set_document_margins(doc)

    styles = doc.styles
    if "Normal" in styles:
        normal_style = get_paragraph_style(doc, "Normal")
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)

    if "Heading 1" in styles:
        heading1_style = get_paragraph_style(doc, "Heading 1")
        heading1_style.font.name = "Times New Roman"
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True

    add_title(doc, job.visible_title)
    add_week_display(doc, job.week_display)
    add_opening_blockquote(doc, job.subtitle)
    add_markdown_body_minimal(doc, source_body)

    logger.debug("Saving DOCX to %s", output_path)
    doc.save(str(output_path))
    logger.info("Wrote DOCX: %s", output_path)


def build_job(source_md: Path, step3_weeks_dir: Path) -> ChapterJob:
    week_num = extract_week_num_from_filename(source_md)
    logger.debug("Matched source chapter %s to week %s", source_md.name, week_num)

    metadata_path = step3_weeks_dir / f"Week {week_num}" / DEFAULT_METADATA_NAME_TEMPLATE.format(week=week_num)

    if not metadata_path.exists():
        raise FileNotFoundError(f"Metadata file not found for week {week_num}: {metadata_path}")

    metadata = load_json(metadata_path)
    logger.debug("Loaded metadata from %s", metadata_path)

    title = metadata.get("title", "").strip()
    subtitle = metadata.get("subtitle", "").strip()
    period_label = metadata.get("period_label", "").strip()

    window = metadata.get("window", {}) or {}
    start_date = str(window.get("start_date", "")).strip()
    end_date = str(window.get("end_date", "")).strip()

    if not title:
        raise ValueError(f"Missing title in {metadata_path}")
    if not subtitle:
        raise ValueError(f"Missing subtitle in {metadata_path}")

    slug = slugify_phrase(title, max_words=5)
    visible_title = visible_title_from_metadata(title)
    week_display = week_display_from_metadata(week_num, period_label, start_date, end_date)
    logger.debug("Built ChapterJob for week %03d with slug '%s'", week_num, slug)

    return ChapterJob(
        week_num=week_num,
        source_md=source_md,
        source_metadata=metadata_path,
        title=title,
        visible_title=visible_title,
        subtitle=subtitle,
        period_label=period_label,
        start_date=start_date,
        end_date=end_date,
        slug=slug,
        week_display=week_display,
    )



def iter_source_files(
    chapters_dir: Path,
    weeks: Optional[set[int]] = None,
) -> Iterable[Path]:
    files = sorted(chapters_dir.glob(DEFAULT_CHAPTER_GLOB))
    for path in files:
        n = extract_week_num_from_filename(path)
        if weeks is not None and n not in weeks:
            continue
        yield path


# Helper: find any existing outputs for a week
def find_existing_week_outputs(output_dir: Path, week_num: int) -> list[Path]:
    patterns = [
        f"week{week_num:03d}_*.md",
        f"week{week_num:03d}_*.docx",
    ]
    matches: list[Path] = []
    for pattern in patterns:
        matches.extend(sorted(output_dir.glob(pattern)))
    return matches


# Helper to format output-path lists cleanly for logs and user messages.
def format_path_list(paths: list[Path]) -> str:
    return ", ".join(str(path) for path in paths)


def get_script_dir() -> Path:
    return Path(__file__).resolve().parent


def default_chapters_dir() -> Path:
    return get_script_dir().parent / "08_RewrittenChapters"


def default_output_dir() -> Path:
    return get_script_dir().parent / "09_VellumImportDocuments"


def default_step3_weeks_dir() -> Path:
    return get_script_dir().parent.parent / "Step 3" / "Weeks"


def configure_logging(level_name: str) -> None:
    level = getattr(logging, level_name.upper(), logging.INFO)
    logging.basicConfig(
        level=level,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )


# -----------------------------
# Main
# -----------------------------

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Build Vellum import markdown/docx files from rewritten chapter markdown + Step 3 metadata."
    )
    parser.add_argument(
        "--chapters-dir",
        default=str(default_chapters_dir()),
        help="Path to Trump Diary/08_RewrittenChapters (default: inferred from script location)",
    )
    parser.add_argument(
        "--step3-weeks-dir",
        default=str(default_step3_weeks_dir()),
        help="Path to Step 3/Weeks (default: inferred from script location)",
    )
    parser.add_argument(
        "--output-dir",
        default=str(default_output_dir()),
        help="Path to Trump Diary/09_VellumImportDocuments (default: inferred from script location)",
    )
    parser.add_argument(
        "--week",
        type=int,
        required=True,
        help="First week to build",
    )
    parser.add_argument(
        "--weeks",
        type=int,
        default=1,
        help="Number of consecutive weeks to build starting from --week (default: 1)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show planned outputs without writing files",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite existing output files",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        help="Logging verbosity (default: INFO)",
    )

    args = parser.parse_args()

    if args.weeks < 1:
        parser.error("--weeks must be at least 1.")

    configure_logging(args.log_level)

    chapters_dir = Path(args.chapters_dir).expanduser().resolve()
    step3_weeks_dir = Path(args.step3_weeks_dir).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    weeks_filter = set(range(args.week, args.week + args.weeks))

    logger.info("Starting Vellum import build")
    logger.info("Chapters dir: %s", chapters_dir)
    logger.info("Step 3 Weeks dir: %s", step3_weeks_dir)
    logger.info("Output dir: %s", output_dir)
    logger.info("Start week: %s", args.week)
    logger.info("Weeks to build: %s", args.weeks)
    logger.info("Computed week set: %s", sorted(weeks_filter))
    logger.info("Dry run: %s", args.dry_run)
    logger.info("Force overwrite: %s", args.force)

    if not chapters_dir.exists():
        print(f"ERROR: chapters directory does not exist: {chapters_dir}", file=sys.stderr)
        return 1
    if not step3_weeks_dir.exists():
        print(f"ERROR: Step 3 Weeks directory does not exist: {step3_weeks_dir}", file=sys.stderr)
        return 1

    if not args.dry_run:
        output_dir.mkdir(parents=True, exist_ok=True)

    source_files = list(iter_source_files(chapters_dir, weeks=weeks_filter))
    if not source_files:
        logger.warning(
            "No source chapter files matched pattern '%s' in %s for requested weeks %s",
            DEFAULT_CHAPTER_GLOB,
            chapters_dir,
            sorted(weeks_filter),
        )
        return 0

    logger.info("Found %d source chapter(s) to process", len(source_files))

    built = 0
    skipped = 0

    for index, source_md in enumerate(source_files, start=1):
        logger.info("Processing chapter %d of %d: %s", index, len(source_files), source_md.name)
        logger.debug("Week %03d selected for processing", extract_week_num_from_filename(source_md))
        try:
            job = build_job(source_md, step3_weeks_dir)
            source_body = read_source_body(source_md, expected_title=job.title)

            base_name = f"week{job.week_num:03d}_{job.slug}"
            md_out = output_dir / f"{base_name}.md"
            docx_out = output_dir / f"{base_name}.docx"
            existing_week_outputs = find_existing_week_outputs(output_dir, job.week_num)
            if existing_week_outputs and not args.force:
                existing_str = format_path_list(existing_week_outputs)
                raise OverwriteRequiredError(
                    f"Output file(s) already exist for week {job.week_num:03d}: {existing_str}. "
                    "Re-run with --force to overwrite."
                )

            logger.info("Title: %s", job.title)
            logger.debug("Subtitle: %s", job.subtitle)
            logger.debug("Metadata source: %s", job.source_metadata)
            logger.info("Planned outputs: %s, %s", md_out.name, docx_out.name)

            print(f"[week {job.week_num:03d}]")
            print(f"  source md : {source_md}")
            print(f"  metadata  : {job.source_metadata}")
            print(f"  title     : {job.title}")
            print(f"  subtitle  : {job.subtitle}")
            print(f"  outputs   : {md_out.name}, {docx_out.name}")

            if not args.dry_run:
                if args.force:
                    for existing_path in existing_week_outputs:
                        logger.info("Removing existing output before rewrite: %s", existing_path)
                        existing_path.unlink()
                md_text = markdown_for_output(job, source_body)
                logger.debug("Writing enriched markdown to %s", md_out)
                md_out.write_text(md_text, encoding="utf-8")
                logger.info("Wrote markdown: %s", md_out)
                write_docx(job, source_body, docx_out)

            built += 1

        except OverwriteRequiredError as exc:
            skipped += 1
            logger.warning("Skipping %s because overwrite is required", source_md.name)
            print(f"SKIP {source_md.name}: {exc}", file=sys.stderr)
            continue

        except Exception as exc:
            logger.exception("Failed while processing %s", source_md.name)
            print(f"ERROR processing {source_md.name}: {exc}", file=sys.stderr)
            return 1

    logger.info("Build complete. Built %d chapter(s); skipped %d chapter(s).", built, skipped)
    print(f"Done. Built {built} chapter(s); skipped {skipped} chapter(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())