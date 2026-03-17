#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from json import JSONDecodeError
import logging
import sys
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore

try:
    import jsonschema  # type: ignore
except ImportError:  # pragma: no cover
    jsonschema = None  # type: ignore


DEFAULT_SCHEMA_FILE = Path(
    "/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/11_Schemas/phase2_year_themes_chapter.schema.json"
)
DEFAULT_PACKAGE_FILE = Path(
    "/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/14_Data/year_themes_chapter_package_v1.json"
)
DEFAULT_PROMPT_FILE = Path(
    "/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/12_Prompts/phase2_year_themes_chapter_prompt_v1.md"
)
DEFAULT_OUTPUT_DIR = Path(
    "/Volumes/PRINTIFY24/Democracy Clock Automation/Trump Diary/10_BackMatter"
)

SECTION_STEM = "year_themes_chapter"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build the Trump Chronicles year-themes chapter package, markdown scaffold, and DOCX scaffold."
    )
    parser.add_argument(
        "--package-file",
        type=Path,
        default=DEFAULT_PACKAGE_FILE,
        help="Path to the year themes chapter package JSON.",
    )
    parser.add_argument(
        "--schema-file",
        type=Path,
        default=DEFAULT_SCHEMA_FILE,
        help="Optional schema file for validation.",
    )
    parser.add_argument(
        "--prompt-file",
        type=Path,
        default=DEFAULT_PROMPT_FILE,
        help="Optional prompt file to record in notes if present.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help="Directory for output artifacts.",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Overwrite existing outputs.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show planned work without writing files.",
    )
    parser.add_argument(
        "--created-by",
        type=str,
        default="build_year_themes_chapter.py",
        help="Value for build.created_by in the normalized JSON output.",
    )
    parser.add_argument(
        "--prompt-version",
        type=str,
        default="year-themes-chapter-v1",
        help="Value for build.prompt_version in the normalized JSON output.",
    )
    parser.add_argument(
        "--model",
        type=str,
        default=None,
        help="Optional model identifier for provenance.",
    )
    parser.add_argument(
        "--git-commit",
        type=str,
        default=None,
        help="Optional git commit for provenance.",
    )
    parser.add_argument(
        "--log-level",
        "--level",
        dest="log_level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="Logging level.",
    )
    return parser.parse_args()


def configure_logging(level: str) -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper()),
        format="%(asctime)s [%(levelname)s] %(message)s",
    )


def load_json(path: Path) -> Dict[str, Any]:
    logging.debug("Loading JSON file: %s", path)
    try:
        with path.open("r", encoding="utf-8") as fh:
            payload = json.load(fh)
    except JSONDecodeError as exc:
        raise ValueError(
            f"{path}: invalid JSON at line {exc.lineno}, column {exc.colno} (char {exc.pos}): {exc.msg}"
        ) from exc
    if not isinstance(payload, dict):
        raise ValueError(f"{path}: expected a JSON object")
    logging.debug("Loaded JSON object from %s with keys: %s", path, ", ".join(sorted(payload.keys())))
    return payload


def validate_package(package: Dict[str, Any], schema_path: Path) -> None:
    if not schema_path.exists():
        logging.warning("Schema file not found, skipping validation: %s", schema_path)
        return
    if jsonschema is None:
        logging.warning("jsonschema is not installed, skipping validation against %s", schema_path)
        return
    logging.debug("Loading schema for validation: %s", schema_path)
    schema = load_json(schema_path)
    jsonschema.validate(instance=package, schema=schema)
    logging.info("Validated package against schema: %s", schema_path)



def ensure_writable(output_dir: Path, force: bool) -> Dict[str, Path]:
    paths = {
        "json": output_dir / f"{SECTION_STEM}.json",
        "md": output_dir / f"{SECTION_STEM}.md",
        "docx": output_dir / f"{SECTION_STEM}.docx",
    }
    existing = [p for p in paths.values() if p.exists()]
    if existing and not force:
        raise FileExistsError(
            "Output file(s) already exist: " + ", ".join(str(p) for p in existing) + ". Re-run with --force."
        )
    return paths



def normalize_package(
    package: Dict[str, Any],
    created_by: str,
    prompt_version: str,
    model: Optional[str],
    git_commit: Optional[str],
    prompt_file: Path,
) -> Dict[str, Any]:
    payload = json.loads(json.dumps(package))
    now_utc = datetime.now(UTC)
    payload.setdefault("build", {})
    payload["build"]["created_at"] = now_utc.replace(microsecond=0).isoformat().replace("+00:00", "Z")
    payload["build"]["created_by"] = created_by
    payload["build"]["model"] = model
    payload["build"]["run_id"] = f"year-themes-{now_utc.strftime('%Y%m%dT%H%M%SZ')}-{uuid.uuid4().hex[:8]}"
    payload["build"]["prompt_version"] = prompt_version
    payload["build"]["git_commit"] = git_commit

    notes = payload.setdefault("notes", [])
    if isinstance(notes, list) and prompt_file.exists():
        note = f"Prompt file: {prompt_file}"
        if note not in notes:
            notes.append(note)

    return payload



def render_markdown(package: Dict[str, Any]) -> str:
    lines: List[str] = []
    title = package["chapter_title"]
    subtitle = package.get("chapter_subtitle")
    chapter_brief = package["chapter_brief"]
    closing_brief = package["closing_brief"]

    lines.append(f"# {title}")
    lines.append("")
    if subtitle:
        lines.append(f"> {subtitle}")
        lines.append("")

    lines.append(chapter_brief["opening_brief"])
    lines.append("")

    for theme in sorted(package["themes"], key=lambda t: t["sort_order"]):
        lines.append(f"## {theme['theme_name']}")
        lines.append("")
        lines.append(theme["framing_line"])
        lines.append("")
        lines.append(theme["section_brief"])
        lines.append("")

        recurring_forms = "; ".join(form["form_name"] for form in theme.get("recurring_forms", []))
        if recurring_forms:
            lines.append(f"**Recurring forms:** {recurring_forms}.")
            lines.append("")

        illustrative = ", ".join(f"Week {item['week']}" for item in theme.get("illustrative_weeks", []))
        if illustrative:
            lines.append(f"**Illustrative weeks for drafting focus:** {illustrative}.")
            lines.append("")

        lines.append(theme["bridge_forward"])
        lines.append("")

    lines.append(closing_brief["closing_brief"])
    lines.append("")
    return "\n".join(lines).rstrip() + "\n"



def render_docx(markdown_text: str, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()
    for block in markdown_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            p = doc.add_paragraph()
            p.style = "Heading 1"
            p.add_run(block[2:].strip())
            continue
        if block.startswith("> "):
            p = doc.add_paragraph()
            p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Quote"
            p.add_run(block[2:].strip())
            continue
        if block.startswith("## "):
            p = doc.add_paragraph()
            p.style = "Heading 2"
            p.add_run(block[3:].strip())
            continue
        if block.startswith("**Recurring forms:**"):
            p = doc.add_paragraph()
            run = p.add_run("Recurring forms: ")
            run.bold = True
            p.add_run(block[len("**Recurring forms:**"):].strip())
            continue
        if block.startswith("**Illustrative weeks for drafting focus:**"):
            p = doc.add_paragraph()
            run = p.add_run("Illustrative weeks for drafting focus: ")
            run.bold = True
            p.add_run(block[len("**Illustrative weeks for drafting focus:**"):].strip())
            continue
        doc.add_paragraph(block)
    doc.save(str(output_path))



def main() -> int:
    args = parse_args()
    configure_logging(args.log_level)

    logging.info("Building year-themes chapter artifacts")
    logging.info("Package file: %s", args.package_file)
    logging.info("Schema file: %s", args.schema_file)
    logging.info("Output dir: %s", args.output_dir)

    if args.dry_run:
        logging.info("[DRY RUN] Would read package: %s", args.package_file)
        logging.info("[DRY RUN] Would validate against schema: %s", args.schema_file)
        logging.info("[DRY RUN] Would use prompt file: %s", args.prompt_file)
        for suffix in ("json", "md", "docx"):
            logging.info("[DRY RUN] Would write: %s", args.output_dir / f"{SECTION_STEM}.{suffix}")
        return 0

    try:
        args.output_dir.mkdir(parents=True, exist_ok=True)
        output_paths = ensure_writable(args.output_dir, args.force)
        logging.info("Loading package JSON")
        package = load_json(args.package_file)
        logging.info("Validating package against schema")
        validate_package(package, args.schema_file)
        logging.info("Normalizing package build metadata")
        normalized = normalize_package(
            package=package,
            created_by=args.created_by,
            prompt_version=args.prompt_version,
            model=args.model,
            git_commit=args.git_commit,
            prompt_file=args.prompt_file,
        )
        logging.info("Rendering markdown scaffold")
        markdown = render_markdown(normalized)
    except Exception as exc:
        logging.error("Build setup failed: %s", exc)
        return 1

    try:
        with output_paths["json"].open("w", encoding="utf-8") as fh:
            json.dump(normalized, fh, indent=2, ensure_ascii=False)
            fh.write("\n")
        with output_paths["md"].open("w", encoding="utf-8") as fh:
            fh.write(markdown)
        render_docx(markdown, output_paths["docx"])
    except Exception as exc:
        logging.error("Failed writing outputs: %s", exc)
        return 1

    logging.info("WROTE: %s", output_paths["json"])
    logging.info("WROTE: %s", output_paths["md"])
    logging.info("WROTE: %s", output_paths["docx"])
    return 0


if __name__ == "__main__":
    sys.exit(main())
